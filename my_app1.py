from docx import Document
from pathlib import Path
from PIL import Image
from docx.shared import Inches
import pyttsx3

# example how python text to speech works   
# pyttsx3.speak('Menifest')

# fun def which converts text to speech
def speak(text):
    pyttsx3.speak(text)

# fun def which convert image type to JPG
def image_to_jpg(image_path):
    path = Path(image_path)
    if path.suffix not in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
        jpg_image_path = f'{path.parent / path.stem}_result.jpg'
        Image.open(image_path).convert('RGB').save(jpg_image_path)
        return jpg_image_path
    return image_path

document = Document()

# 1st version name phone and email hardcoded
# name = 'Dilip'
# phone_number = '0008884445'
# email = 'dilip@python.com'
# or
# 2nd version using cmd to take input details
name = input('What is your name? ')
speak('Hello '+name+' hope you are doing good today')

speak('Could you please enter you phone number?')
phone_number = input('what is your phone number? ')
speak('cheers! Your mobile number is '+ phone_number)

speak('could you please provide your personal email address?')
email = input('what is your eamil address? ')
speak('cheers!')

# profile picture
document.add_picture(image_to_jpg('dilip.jfif'), width=Inches(1.5))

# details(name, phone and email)
document.add_paragraph(name + ' | ' + phone_number+ ' | ' +email)

# About me
document.add_heading('About me')

# 1st version of about me hard coded
# document.add_paragraph('I am full stack developer and started loving my job of coding and doing new stuffs.')
# or
# 2nd version buy taking input via cmd
speak('Could you please brief youself?')
document.add_paragraph(input('Tell me about yourself? '))

# Work experience
document.add_heading('Work Experience')
speak('Provide your past company working details!')
p = document.add_paragraph()
company = input('What is your company? ')
from_date = input('started working from(date-DD/MM/YY)? ')
end_date = input('end date? ')
p.add_run(company + ' ').bold = True
p.add_run('('+from_date+'-'+end_date+')\n').italic=True
experince_det = input('Describe you experince at '+ company+ '? ')
p.add_run(experince_det)

# more than one company
while True:
    has_more_experience = input('Do you have more experience? Yes or NO ')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()
        company = input('What is your company? ')
        from_date = input('started working from(date-DD/MM/YY)? ')
        end_date = input('end date? ')
        p.add_run(company + ' ').bold = True
        p.add_run('('+from_date+'-'+end_date+')\n').italic=True
        experince_det = input('Describe you experince at '+ company+ '? ')
        p.add_run(experince_det)
    else:
        break

# Skills
document.add_heading('Skills')
speak('Enter your primary skill')
skill = input('Enter your primary skill : ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# loop for secondary skills
while True:
    has_any_secondary_skills = input('Do you have any secondary skills? Yes or No ')
    if has_any_secondary_skills.lower() == 'yes':
        skill = input('Enter your secondary skill : ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Python programng on second day of learning python@2022'

document.save('cv.docx')